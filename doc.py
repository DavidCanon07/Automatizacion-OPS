from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

doc = Document()

# Título principal
title = doc.add_heading('🗝️ VALIDADOR 🗝️', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_heading('MANUAL FUNCIONAL Y TÉCNICO', level=1)
doc.add_heading('Sistema Pre-Validador y Consolidador de Archivos OPS', level=2)

# Tabla de metadatos
table = doc.add_table(rows=4, cols=2)
table.style = 'Table Grid'
table.cell(0, 0).text = 'Versión'
table.cell(0, 1).text = '2.1'
table.cell(1, 0).text = 'Fecha'
table.cell(1, 1).text = datetime.now().strftime('%d/%m/%Y')
table.cell(2, 0).text = 'Autor'
table.cell(2, 1).text = 'David Cañon'
table.cell(3, 0).text = 'Clasificación'
table.cell(3, 1).text = 'Interno'

# ========== 1. INTRODUCCIÓN ==========
doc.add_heading('1. Introducción', level=1)
doc.add_paragraph("""El presente documento describe el funcionamiento del sistema 🗝️ VALIDADOR 🗝️, desarrollado como solución integral de automatización para el proceso OPS (Órdenes de Pago y Servicios). El sistema cubre dos grandes frentes: la pre-validación de la calidad y estructura de los archivos antes de su carga en el entorno productivo, y la consolidación y generación de los entregables finales del proceso.

El sistema está diseñado para ejecutarse en entornos Windows con Python 3.9+ y Microsoft Excel instalado (requerido para algunas funciones de consolidación). Se lanza mediante el archivo executer.bat o directamente con el comando python validador.py.""")

# ========== 2. OBJETIVO ==========
doc.add_heading('2. Objetivo del Sistema', level=1)
objetivos = [
    "Validar la estructura de columnas obligatorias en los archivos fuente.",
    "Verificar la longitud máxima y mínima de campos críticos.",
    "Detectar valores inválidos en campos de control (tipo, filler, entidad, justificación contable).",
    "Identificar campos vacíos y registros duplicados.",
    "Consolidar múltiples archivos OPS en un único archivo unificado.",
    "Generar el archivo en el formato de plantilla OPS requerido por el proceso.",
    "Consolidar el archivo de débitos LATAM/SODIMAC/FALABELLA/MERCADO PAGO.",
    "Poblar automáticamente la Solicitud OPS con los valores calculados.",
    "Exportar el archivo plano (.txt) para carga en sistema.",
    "Comprimir los entregables finales en formato .zip."
]
for obj in objetivos:
    doc.add_paragraph(obj, style='List Bullet')

# ========== 3. ARQUITECTURA ==========
doc.add_heading('3. Arquitectura del Sistema', level=1)
arch_table = doc.add_table(rows=7, cols=2)
arch_table.style = 'Table Grid'
arch_data = [
    ('Módulo / Paquete', 'Descripción'),
    ('validador.py', 'Orquestador principal. Controla el menú interactivo y el flujo completo del proceso.'),
    ('Consolidacion/', 'Paquete de consolidación. Contiene los módulos de unificación de OPS, débitos, solicitud, carga de estructura y compresión.'),
    ('validacion/', 'Paquete de pre-validación. Contiene los módulos de carga, validaciones, reportes y utilidades.'),
    ('Configuracion_parametros.py', 'Archivo centralizado de parámetros, rutas, mapeos y reglas de negocio.'),
    ('preguntas.py', 'Módulo de validación interactiva de requisitos mínimos antes de iniciar el proceso.'),
    ('executer.bat', 'Script de arranque del sistema en entornos Windows.')
]
for i, (mod, desc) in enumerate(arch_data):
    arch_table.cell(i, 0).text = mod
    arch_table.cell(i, 1).text = desc

# ========== 4. DESCRIPCIÓN DE MÓDULOS ==========
doc.add_heading('4. Descripción de Módulos', level=1)

# 4.1 validador.py
doc.add_heading('4.1 validador.py — Orquestador Principal', level=2)
doc.add_paragraph("""Es el punto de entrada del sistema. Presenta al usuario un menú interactivo de 8 opciones y coordina la ejecución de todos los módulos de validación y consolidación. Maneja el ciclo completo de vida de una ejecución: inicialización, limpieza, validación, consolidación y generación de entregables.

Responsabilidades principales:
- Verificar y crear la estructura de carpetas requerida para el proceso.
- Validar la existencia de las plantillas base (estructuras) y copiarlas desde la ruta de instalación si no existen.
- Limpiar archivos temporales de ejecuciones anteriores antes de iniciar.
- Coordinar la ejecución secuencial de todas las validaciones (pasos 01 al 09 más duplicados).
- Registrar logs de éxito y error en archivos .txt con marca de tiempo.
- Controlar el flujo de continuación/parada según el tipo de error detectado.""")

doc.add_heading('Menú de opciones:', level=3)
menu_table = doc.add_table(rows=9, cols=2)
menu_table.style = 'Table Grid'
menu_data = [
    ('Opción', 'Acción'),
    ('1', 'Consolidar archivos OPS y ejecutar todas las validaciones (pasos 01–09 + duplicados).'),
    ('2', 'Exportar el DataFrame consolidado como archivo temporal unificado (temp_archivos_unificados.xlsx).'),
    ('3', 'Montar el archivo \'Formato OPS DDMMYYYY.xlsx\' copiando la plantilla limpia y cargando los datos.'),
    ('4', 'Consolidar el archivo de débitos \'Detalle_LATAM_SODIMAC_FALABELLA_MERCADO PAGO_DDMMYYYY.xlsx\'.'),
    ('5', 'Llenar automáticamente la \'Solicitud OPS DDMMYYYY.xlsx\' con valores calculados del formato OPS.'),
    ('6', 'Generar el archivo plano \'OPS DDMMYYYY.txt\' para carga en sistema.'),
    ('7', 'Comprimir el archivo \'Formato OPS DDMMYYYY.xlsx\' en un .zip.'),
    ('8', 'Salir del programa.')
]
for i, (op, acc) in enumerate(menu_data):
    menu_table.cell(i, 0).text = op
    menu_table.cell(i, 1).text = acc

# 4.2 Paquete Consolidacion/
doc.add_heading('4.2 Paquete Consolidacion/', level=2)

# Unificado.py
doc.add_heading('4.2.1 Unificado.py', level=3)
doc.add_paragraph("""
limpiar_numero_latino(valor) — Convierte valores numéricos en formato latino (1.234,56) a float Python (1234.56). Detecta automáticamente si el separador de miles es punto o coma, y maneja múltiples combinaciones de formatos (con/sin miles, solo coma, solo punto). Retorna None si el valor está vacío y el valor original si no pudo convertir.

exportar_excel(df, ruta, sheet_name) — Selecciona las columnas requeridas para el entregable final, aplica la conversión numérica sobre la columna 'valor ajuste', agrega una columna de índice correlativo (N°) y exporta el DataFrame a Excel usando openpyxl. Muestra estadísticas de conversión en consola.
""")

# Debitos.py
doc.add_heading('4.2.2 Debitos.py', level=3)
doc.add_paragraph("""
obtener_debitos(carpeta_archivos, clave_debitos, exts) — Recorre recursivamente la carpeta de archivos buscando archivos cuyo nombre contenga la clave_debitos (por defecto 'Detalle_LATAM'). Por cada archivo encontrado, lee el Excel con pandas, formatea las columnas de fecha ('Fecha Canje', 'Fecha Comprobante') al formato DD/MM/YYYY, limpia y convierte las columnas numéricas de valor a float, y acumula los DataFrames. Retorna el DataFrame consolidado de todos los archivos de débitos encontrados.

exportar_debitos(df_total_debitos, ruta, sheet_name) — Exporta el DataFrame consolidado de débitos a un archivo Excel en la ruta especificada usando openpyxl.
""")

# Solicitud_ops.py
doc.add_heading('4.2.3 Solicitud_ops.py', level=3)
doc.add_paragraph("""
obtener_solicitud_ops(...) — Abre el archivo de Formato OPS con Excel COM (win32com) para forzar el recálculo de fórmulas, guarda los valores calculados y luego los lee con openpyxl (data_only=True). Copia los valores de las celdas origen a las celdas destino de la Solicitud OPS según el MAPEO_SOLICITUD, y escribe la fecha actual en D11.

extraer_cuentas_y_descripciones(...) — Usando Excel COM, recorre las columnas J (Descripción), K (Cuenta contable) y N (Código MIR) del Formato OPS desde la fila de inicio hasta la última fila con datos. Deduplica los valores de cada columna y los concatena con un separador configurable. Escribe los resultados concatenados en las celdas destino de la Solicitud OPS.

exportar_txt_limpio(ruta_formato_ops, hoja, rango, archivo_txt) — Lee el Formato OPS desde la fila 9 (skiprows=8), elimina filas vacías y filas donde el primer campo es cero o vacío, y exporta cada fila como una línea del archivo .txt concatenando todos los valores sin separador. Evita dejar salto de línea al final del archivo.
""")

# carga_estructura.py
doc.add_heading('4.2.4 carga_estructura.py', level=3)
doc.add_paragraph("""
cargar_estructura(ruta_archivo_unificado, ruta_libro_base, hoja_base, MAPEO, fila_encabezados) — Función genérica de carga de datos desde un archivo origen a una plantilla destino. Lee los encabezados del origen (siempre fila 1) y del destino (fila configurable, por defecto 8). Valida el mapeo columna a columna, identifica la primera fila libre en el destino y copia los datos respetando el mapeo. Crea el archivo destino si no existe.
""")

# comprimir.py
doc.add_heading('4.2.5 comprimir.py', level=3)
doc.add_paragraph("""
comprimir_excel(archivo_excel, carpeta_destino) — Comprime un archivo Excel en formato .zip con compresión DEFLATE. Genera el .zip en la carpeta destino especificada con el mismo nombre base del archivo de entrada.
""")

# 4.3 Paquete validacion/
doc.add_heading('4.3 Paquete validacion/', level=2)

# carga.py
doc.add_heading('4.3.1 carga.py', level=3)
doc.add_paragraph("""
validar_columnas(df, archivo) — Verifica que el DataFrame contenga exactamente las columnas esperadas (definidas en Campos_a_validar) hasta el límite de la columna 'DIGITO DE VERIFICACION'. Lanza excepción crítica si hay columnas faltantes. Lanza excepción de advertencia si hay columnas extra, convirtiendo los nombres tipo 'Unnamed: N' al formato legible 'Columna X'.

recortar_footer_dinamico(df) — Elimina dinámicamente las filas del footer de los archivos Excel. Evalúa las columnas ancla definidas en Configuracion_parametros.py y considera válida una fila si tiene dato en al menos 3 de ellas simultáneamente. Descarta todo lo que viene después de la última fila válida, sin depender de un número fijo de filas de footer.

obtener_datos(carpeta_archivos, clave, exts) — Recorre recursivamente la carpeta de archivos buscando archivos que contengan la clave OPS en su nombre y tengan extensión .xlsx o .xls. Por cada archivo: lee con skiprows=6, aplica recortar_footer_dinamico(), valida columnas, agrega la fecha actual y el nombre de la carpeta como columnas FECHA y Proceso.

concatenar_datos(carpeta_archivos, clave, exts) — Llama a obtener_datos() y consolida todos los DataFrames en uno solo con pd.concat().
""")

# validaciones.py
doc.add_heading('4.3.2 validaciones.py', level=3)
val_table = doc.add_table(rows=10, cols=3)
val_table.style = 'Table Grid'
val_data = [
    ('Función', 'Tipo', 'Descripción'),
    ('validar_largo_campos(df)', 'ERROR', 'Valida longitud exacta o rango (min, max) por campo según largo_campos. Genera 02 - errores_largo_campos.xlsx.'),
    ('validar_campos_vacios(df)', 'ALERTA', 'Detecta campos vacíos en columnas obligatorias. Genera 04 - alertas_campos.xlsx pero no detiene el proceso.'),
    ('validar_columna_tipo(df)', 'ERROR', 'Valida que la columna \'tipo\' solo contenga \'P\' o \'N\' en mayúscula. Genera 03 - errores_columna_tipo.xlsx.'),
    ('validar_redondeo_valores(df)', 'ERROR', 'Valida que \'valor ajuste\' tenga exactamente 2 decimales. Genera 05 - errores_redondeo.xlsx.'),
    ('validar_inicio_numero_cuenta(df, campo, prefijo)', 'ERROR', 'Valida que \'numero de la cuenta\' inicie por \'1\' (ahorros) o \'2\' (corriente). Genera 06 - errores_inicio_numero_cuenta.xlsx.'),
    ('validar_entidad_cuenta(df)', 'ERROR', 'Valida que \'Entidad de la cuenta\' sea siempre \'0013\'. Genera 07 - errores_entidad_cuenta.xlsx.'),
    ('validar_filler(df)', 'ERROR', 'Valida que el campo \'filler\' sea siempre \'0\'. Genera 08 - errores_filler.xlsx.'),
    ('validar_justificacion_contable(df)', 'ERROR', 'Valida que la justificación contable corresponda a los valores permitidos según el tipo (P o N). Genera 09 - errores_justificacion.xlsx.'),
    ('validar_duplicados(df)', 'ALERTA', 'Detecta registros duplicados por combinación de campos clave. Genera ALERTA_DUPLICADOS.xlsx. No detiene el proceso.')
]
for i, row in enumerate(val_data):
    for j, cell_val in enumerate(row):
        val_table.cell(i, j).text = cell_val

# reportes.py
doc.add_heading('4.3.3 reportes.py', level=3)
doc.add_paragraph("""
borrar_archivos_temporales() — Elimina todos los archivos de reportes, logs y temporales generados en ejecuciones anteriores para garantizar un entorno limpio al inicio de cada ejecución.

borrar_archivo_carpeta_formato_ops(formatos) — Elimina todos los archivos .xlsx de la carpeta de entregables finales (Archivos OPS/).

borrar_carpeta_comprimido(formatos) — Elimina todos los archivos .zip de la carpeta de entregables finales.

exportar_errores(df_errores, ruta, mensaje, sheet_name) — Función generalizada para exportar DataFrames de errores a Excel. Ajusta el índice al número de fila real de Excel (+8 por los skiprows), ordena por fila, crea la carpeta si no existe y lanza una excepción con el mensaje indicado para detener el proceso.
""")

# utils.py
doc.add_heading('4.3.4 utils.py', level=3)
doc.add_paragraph("""
formatea_nombre_columna(nombre) — Convierte nombres de columnas tipo 'Unnamed: N' al formato legible 'Columna X' usando la letra de columna Excel correspondiente.

escribir(texto, velocidad) — Imprime texto con efecto de máquina de escribir para mejorar la experiencia de usuario en consola.

input_con_efecto(texto, velocidad) — Muestra texto con efecto de máquina de escribir y retorna el input del usuario.

crear_carpeta_si_no_existe(...) — Crea la estructura de carpetas necesaria para el proceso si alguna no existe (C:\\validador, carpetas_OPS, Control de ejecuciones, Archivos OPS, Estructuras).

validar_estructuras(...) — Verifica que las tres plantillas base (Estructura_Formato_OPS.xlsx, Estructura_Debitos.xlsx, Estructura_Solicitud_OPS.xlsx) existan en la carpeta Estructuras. Si alguna falta, la copia desde la ruta de instalación C:\\Programa_validador\\estructuras_base.
""")

# 4.4 Configuracion_parametros.py
doc.add_heading('4.4 Configuracion_parametros.py — Centro de Control', level=2)
doc.add_paragraph("""Centraliza todas las reglas de negocio, rutas y configuraciones del sistema. Modificar este archivo permite adaptar el validador a cambios del proceso sin alterar la lógica de los demás módulos.""")

param_table = doc.add_table(rows=11, cols=2)
param_table.style = 'Table Grid'
param_data = [
    ('Parámetro', 'Descripción'),
    ('columna_ancla', 'Lista de columnas usadas por recortar_footer_dinamico() para detectar filas válidas. Default: [\'Entidad de la cuenta\', \'Centro cuenta\', \'numero de la cuenta\', \'tipo\'].'),
    ('Campos_a_validar', 'Lista completa de columnas esperadas en el archivo OPS, incluidos los dos campos de fecha/proceso implícitos.'),
    ('largo_campos', 'Diccionario con reglas de longitud por campo: int para longitud exacta, tuple (min, max) para rango.'),
    ('justificacion_contable', 'Diccionario con listas de justificaciones válidas para cada tipo de ajuste (\'P\' y \'N\').'),
    ('MAPEO', 'Mapeo columna_origen → columna_destino para la carga al Formato OPS.'),
    ('MAPEO_DEBITOS', 'Mapeo columna_origen → columna_destino para la carga al archivo de débitos.'),
    ('MAPEO_SOLICITUD', 'Mapeo celda_origen → celda_destino para la población de la Solicitud OPS.'),
    ('Rutas (ruta_*)', 'Variables con rutas absolutas a todos los archivos de entrada, salida, errores y logs del sistema.'),
    ('clave / clave_debitos', 'Palabras clave para filtrar archivos OPS (\'OPS\') y de débitos (\'Detalle_LATAM\') por nombre.')
]
for i, row in enumerate(param_data):
    param_table.cell(i, 0).text = row[0]
    param_table.cell(i, 1).text = row[1]

# ========== 5. FLUJO DEL PROCESO ==========
doc.add_heading('5. Flujo del Proceso', level=1)

doc.add_heading('5.1 Flujo de Inicialización (automático al arrancar)', level=2)
doc.add_paragraph("""
1. Arranque del sistema vía executer.bat o python validador.py.
2. Verificación y creación de carpetas: C:\\validador, carpetas_OPS, Control de ejecuciones, Archivos OPS, Estructuras.
3. Validación de plantillas base: se copian desde C:\\Programa_validador\\estructuras_base si no existen.
4. Limpieza de archivos temporales de ejecuciones anteriores.
5. Verificación de requisitos mínimos (preguntas interactivas al usuario).
6. Presentación del menú principal.
""")

doc.add_heading('5.2 Flujo de Validación (Opción 1)', level=2)
doc.add_paragraph("""
1. Carga recursiva de archivos OPS desde carpetas_OPS/ (clave 'OPS', extensión .xlsx/.xls).
2. Recorte dinámico del footer por columnas ancla.
3. Validación de columnas obligatorias (error crítico → detiene).
4. Validación de largo de campos (error crítico → detiene, genera 02-errores.xlsx).
5. Validación de columna 'tipo' — solo 'P' o 'N' (error crítico → detiene, genera 03-errores.xlsx).
6. Validación de campos vacíos (alerta → continúa, genera 04-alertas.xlsx).
7. Validación de redondeo en 'valor ajuste' (error crítico → detiene, genera 05-errores.xlsx).
8. Validación de inicio de 'numero de la cuenta' — prefijos '1' o '2' (error crítico → detiene, genera 06-errores.xlsx).
9. Validación de 'Entidad de la cuenta' — solo '0013' (error crítico → detiene, genera 07-errores.xlsx).
10. Validación de 'filler' — solo '0' (error crítico → detiene, genera 08-errores.xlsx).
11. Validación de 'Justificacion contable' vs. tipo permitido (error crítico → detiene, genera 09-errores.xlsx).
12. Validación de duplicados (alerta → continúa, genera ALERTA_DUPLICADOS.xlsx).
13. Registro en log.txt si todo fue exitoso.
""")

doc.add_heading('5.3 Flujo de Consolidación y Entregables (Opciones 2–7)', level=2)
cons_table = doc.add_table(rows=7, cols=2)
cons_table.style = 'Table Grid'
cons_data = [
    ('Paso', 'Descripción'),
    ('2', 'Exportar el DataFrame validado a temp_archivos_unificados.xlsx. Valida que el archivo fue creado correctamente y es un ZIP válido.'),
    ('3', 'Copiar la plantilla limpia Estructura_Formato_OPS.xlsx a Archivos OPS/FORMATO OPS DDMMYYYY.xlsx. Cargar datos usando el MAPEO de columnas (encabezados en fila 8 de la plantilla).'),
    ('4', 'Consolidar archivos Detalle_LATAM desde carpetas_OPS/. Exportar temporal. Copiar plantilla débitos y cargar datos con MAPEO_DEBITOS (encabezados en fila 1).'),
    ('5', 'Copiar plantilla Solicitud OPS. Usar Excel COM para recalcular fórmulas del Formato OPS. Copiar valores calculados con MAPEO_SOLICITUD. Extraer y concatenar cuentas, descripciones y códigos MIR únicos.'),
    ('6', 'Leer Formato OPS desde fila 9, columna R. Limpiar filas vacías/cero. Exportar como archivo .txt sin salto de línea final.'),
    ('7', 'Comprimir el Formato OPS DDMMYYYY.xlsx en un archivo .zip con el mismo nombre en la carpeta Archivos OPS/.')
]
for i, row in enumerate(cons_data):
    cons_table.cell(i, 0).text = row[0]
    cons_table.cell(i, 1).text = row[1]

# ========== 6. ARCHIVOS GENERADOS ==========
doc.add_heading('6. Archivos Generados', level=1)

doc.add_heading('6.1 Archivos de Validación (C:\\validador\\)', level=2)
# CORREGIDO: Ahora rows=13 (para 12 datos + 1 encabezado = 13 filas)
val_files = doc.add_table(rows=13, cols=2)
val_files.style = 'Table Grid'
val_files_data = [
    ('Archivo', 'Descripción'),
    ('errores.txt', 'Log general de errores críticos con marca de tiempo por cada ejecución.'),
    ('log.txt', 'Log de validaciones exitosas. Se acumula entre ejecuciones para trazabilidad histórica.'),
    ('ejecucion_detallada.log', 'NOVEDAD: Registro detallado de cada paso del flujo con timestamp, duración y cantidad de filas procesadas.'),
    ('02 - errores_largo_campos.xlsx', 'Filas con longitud de campo incorrecta, ordenadas por fila real de Excel.'),
    ('03 - errores_columna_tipo.xlsx', 'Filas con valores inválidos en la columna \'tipo\'.'),
    ('04 - alertas_campos.xlsx', 'Alertas de campos vacíos. No detiene el proceso.'),
    ('05 - errores_redondeo.xlsx', 'Filas con \'valor ajuste\' sin exactamente 2 decimales.'),
    ('06 - errores_inicio_numero_cuenta.xlsx', 'Filas cuyo número de cuenta no inicia por \'1\' o \'2\'.'),
    ('07 - errores_entidad_cuenta.xlsx', 'Filas con entidad de cuenta diferente a \'0013\'.'),
    ('08 - errores_filler.xlsx', 'Filas con filler diferente a \'0\'.'),
    ('09 - errores_justificacion.xlsx', 'Filas con justificación contable no permitida para el tipo indicado.'),
    ('ALERTA_DUPLICADOS.xlsx', 'Registros detectados como duplicados. Alerta — no detiene el proceso.')
]
for i, row in enumerate(val_files_data):
    val_files.cell(i, 0).text = row[0]
    val_files.cell(i, 1).text = row[1]

doc.add_heading('6.2 Archivos Temporales (C:\\validador\\)', level=2)
temp_files = doc.add_table(rows=3, cols=2)
temp_files.style = 'Table Grid'
temp_files_data = [
    ('Archivo', 'Descripción'),
    ('temp_archivos_unificados.xlsx', 'Archivo temporal con el DataFrame consolidado y validado de los archivos OPS. Se elimina al inicio de cada ejecución.'),
    ('temp_archivos_unificados_debitos.xlsx', 'Archivo temporal con el DataFrame consolidado de débitos LATAM. Se elimina al inicio de cada ejecución.')
]
for i, row in enumerate(temp_files_data):
    temp_files.cell(i, 0).text = row[0]
    temp_files.cell(i, 1).text = row[1]

doc.add_heading('6.3 Entregables Finales (C:\\validador\\Archivos OPS\\)', level=2)
final_files = doc.add_table(rows=6, cols=2)
final_files.style = 'Table Grid'
final_files_data = [
    ('Archivo', 'Descripción'),
    ('FORMATO OPS DDMMYYYY.xlsx', 'Archivo principal de OPS en formato de plantilla institucional.'),
    ('FORMATO OPS DDMMYYYY.zip', 'Versión comprimida del Formato OPS para envío.'),
    ('Detalle_LATAM_SODIMAC_FALABELLA_MERCADO PAGO_DDMMYYYY.xlsx', 'Archivo consolidado de débitos LATAM en formato de plantilla.'),
    ('Solicitud OPS DDMMYYYY.xlsx', 'Solicitud OPS poblada automáticamente con valores calculados.'),
    ('OPS DDMMYYYY.txt', 'Archivo plano para carga en sistema. Sin salto de línea final, codificación UTF-8 con BOM.')
]
for i, row in enumerate(final_files_data):
    final_files.cell(i, 0).text = row[0]
    final_files.cell(i, 1).text = row[1]

# ========== 7. MANEJO DE ERRORES ==========
doc.add_heading('7. Manejo de Errores', level=1)
error_table = doc.add_table(rows=4, cols=3)
error_table.style = 'Table Grid'
error_data = [
    ('Tipo', 'Comportamiento', 'Aplica a'),
    ('Error Crítico', 'Detiene el proceso. Registra en errores.txt.', 'Columnas faltantes/extra, longitud de campos, tipo inválido, redondeo, inicio de cuenta, entidad, filler, justificación contable.'),
    ('Alerta', 'Notifica y continúa. Registra en errores.txt.', 'Campos vacíos, registros duplicados.'),
    ('Excepción General', 'Detiene el proceso. Registra en errores.txt.', 'Errores inesperados de ingesta, carga de archivos o estructura de carpetas.')
]
for i, row in enumerate(error_data):
    for j, val in enumerate(row):
        error_table.cell(i, j).text = val

doc.add_paragraph("Todos los errores y alertas incluyen: número de fila real en Excel (índice + 8 por skiprows), nombre del campo involucrado, descripción del error y causal. Los reportes se generan ordenados por fila de Excel para facilitar la corrección.")

# ========== 8. ESTRUCTURA DE CARPETAS Y REQUISITOS ==========
doc.add_heading('8. Estructura de Carpetas y Requisitos de Instalación', level=1)

doc.add_heading('8.1 Estructura de Carpetas en Ejecución', level=2)
doc.add_paragraph("""
C:\\validador\\
├── carpetas_OPS\\ ← Archivos fuente OPS y Detalle_LATAM
│   ├── PROCESO_A\\ ← Subcarpeta por proceso (nombre = columna Proceso)
│   │   └── OPS_*.xlsx
│   └── CBS_L60\\ ← Corresponsales (nombre obligatorio)
│       └── OPS_*.xlsx
├── Archivos OPS\\ ← Entregables finales (generados automáticamente)
├── Control de ejecuciones\\ ← errores.txt, log.txt, ejecucion_detallada.log
├── Estructuras\\ ← Plantillas base (copiadas automáticamente)
└── 02 - errores_*.xlsx ← Reportes de error (raíz de C:\\validador)
""")

doc.add_heading('8.2 Requisitos de Instalación', level=2)
req_table = doc.add_table(rows=7, cols=2)
req_table.style = 'Table Grid'
req_data = [
    ('Componente', 'Detalle'),
    ('Sistema Operativo', 'Windows 10 o superior (requerido para win32com / Excel COM).'),
    ('Python', '3.9 o superior.'),
    ('Microsoft Excel', 'Instalado y licenciado. Requerido para opción 5 (Solicitud OPS) por uso de win32com.'),
    ('Librerías Python', 'pandas, openpyxl, pywin32 (win32com), pathlib (incluida en Python 3.4+).'),
    ('Ruta base', 'C:\\validador (configurable en Configuracion_parametros.py).'),
    ('Ruta instalación', 'C:\\Programa_validador\\estructuras_base (contiene las plantillas base originales).')
]
for i, row in enumerate(req_data):
    req_table.cell(i, 0).text = row[0]
    req_table.cell(i, 1).text = row[1]

doc.add_heading('8.3 Regla de Nombre de Carpeta para Corresponsales', level=2)
doc.add_paragraph("La carpeta que contiene los archivos OPS del proceso de Corresponsales debe llamarse exactamente 'CBS_L60'. El sistema verifica este requisito de forma interactiva antes de iniciar la consolidación. El nombre de la carpeta es usado directamente como valor de la columna Proceso en el archivo consolidado.")

# ========== 9. CATÁLOGO DE JUSTIFICACIONES ==========
doc.add_heading('9. Catálogo de Justificaciones Contables Permitidas', level=1)
doc.add_paragraph("Las siguientes justificaciones son válidas según el tipo de ajuste. La validación es exacta (mayúsculas/minúsculas y acentos incluidos):")

just_table = doc.add_table(rows=16, cols=2)
just_table.style = 'Table Grid'
just_data = [
    ('Tipo N (Cargos / Débitos)', 'Tipo P (Abonos / Créditos)'),
    ('CARGO AJUSTE POR NUEVO MODELO DE RECUPERACIONES', 'ABONO AJUSTE POR NUEVO MODELO DE RECUPERACIONES'),
    ('CARGO DEPOSITOS ELECTRONICOS', 'ABONO DEPOSITOS ELECTRONICOS'),
    ('CARGO DEVOLUCION QR', 'ABONO DEVOLUCION QR'),
    ('CARGO NO APLICADO POR TX INTERNACIONAL CUENTA 400', 'ABONO POR REVERSO PAGO NACIONAL AUTORIZACION'),
    ('CARGO POR DOBLE ABONO A LA TARJETA', 'ABONO RECLAMO ATM'),
    ('CARGO POR TRANSFERENCIA ERRADA', 'ABONO REEMBOLSOS'),
    ('CARGO REEMBOLSOS', 'ABONO RELIQUIDACION COMISIONES'),
    ('CARGO RELIQUIDACION COMISIONES', 'ABONO SALDO A FAVOR'),
    ('CARGO REVERSO PAGO', 'ABONO REVERSO PAGO'),
    ('CARGO SALDO A FAVOR', 'ABONO SALDO A FAVOR ADELANTO DE NOMINA'),
    ('CARGO SALDO A FAVOR ADELANTO DE NOMINA', 'ABONOS'),
    ('CARGOS', 'ABONOS RECLAMACION'),
    ('RECUPERACION TRXS PENDIENTE DE COBRO A CLIENTES, APLICATIVO CONCISO', 'DEVOLUCION INTERN VISA'),
    ('', 'REINTEGRO TRANSFERENCIA P2P SIN COMPENSAR')
]
for i, row in enumerate(just_data):
    just_table.cell(i, 0).text = row[0]
    just_table.cell(i, 1).text = row[1]

# ========== 10. CONTROL DE CAMBIOS ==========
doc.add_heading('10. Control de Cambios', level=1)
control_table = doc.add_table(rows=6, cols=4)
control_table.style = 'Table Grid'
control_data = [
    ('Versión', 'Fecha', 'Autor', 'Descripción del Cambio'),
    ('1.0', '20/02/2026', 'David Cañon', 'Creación inicial del sistema de pre-validación de archivos OPS.'),
    ('1.1', '27/02/2026', 'David Cañon', 'Incorporación de recortar_footer_dinamico() y columna_ancla. Reemplazo de skipfooter fijo por recorte dinámico.'),
    ('1.2', '05/03/2026', 'David Cañon', 'Incorporación de sort_index() en reportes de errores/alertas para ordenar por fila real de Excel. Eliminación de print de depuración.'),
    ('2.0', '05/04/2026', 'David Cañon', 'Incorporación del módulo de Consolidacion (Unificado, Debitos, Solicitud_ops, carga_estructura, comprimir). Menú interactivo de 8 opciones. Generación de archivo plano .txt, compresión .zip y llenado automático de Solicitud OPS.'),
    ('2.1', datetime.now().strftime('%d/%m/%Y'), 'David Cañon', 'Optimización del flujo de validación en Prevalidador. Mejora en recortar_footer_dinamico (mayor flexibilidad). Ampliación de claves de búsqueda en consolidación de débitos. Nuevo log detallado (ejecucion_detallada.log). Soporte para archivos .xls. Corrección de bug en exportación .txt (UTF-8 con BOM). Mejora en menú interactivo (opción de volver).')
]
for i, row in enumerate(control_data):
    for j, val in enumerate(row):
        control_table.cell(i, j).text = str(val)

# Guardar documento
doc.save('Documentacion_OPS_v2_1.docx')
print("✅ Documento generado: Documentacion_OPS_v2_1.docx")